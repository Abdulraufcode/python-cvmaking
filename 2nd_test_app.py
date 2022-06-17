

from docx import Document
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()
#personal information input
name = input("what is you first name? ")
speak("Hello" + name + "How are you... Lets make cv today")
speak("write your second name?")
name_2nd = input("write your secend name? ")
speak("good!....." + name + name_2nd + "now..input..your..hight..and..weight")
age = input("Your hight? ")
weight = input("input your weight.. ")
hight = input("how tall you are? ")
document.add_paragraph(name + " - " + name_2nd + " - " + age + " - " + weight + " - " + hight)
#about me
document.add_heading("About me ")
About_me = input("tell me more about you ")
document.add_paragraph(About_me)
#work experience
document.add_heading("Work Experience")
p = document.add_paragraph()
company = input("company Name? ")
from_date =input("From date ")
to_date = input('To date ')
p.add_run(company + " ").bold = True
p.add_run(from_date + " - " + to_date + "\n").italic = True
experience_details = input("Experience details at " + company)
p.add_run(experience_details)
# more experiences
while True:
    add_more_experience = input("do you have more experiences ? yes or not")
    if add_more_experience == "yes":
        p = document.add_paragraph()
        company = input("company Name? ")
        from_date =input("From date ")
        to_date = input('To date ')
        p.add_run(company + " ").bold = True
        p.add_run(from_date + " - " + to_date + "\n").italic = True
        experience_details = input("Experience details at " + company)
        p.add_run(experience_details)
    else:
        break    
document.add_heading("Skills")
skill = input("input skills ")
p = document.add_paragraph(skill)
p.style = "List Bullet"

while True:
    has_more_skils = input("do you have more skils? yes or no ")
    if has_more_skils == "yes":
        skill = input("input skills ")
        p = document.add_paragraph(skill)
        p.style = "List Bullet"
    else:
        break   

#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "Personal info of Abdul Rauf code"

        



    
    


document.save('personal_info.docx')

